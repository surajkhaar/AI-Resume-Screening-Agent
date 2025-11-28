"""
Resume Parser Module

Extracts structured information from resume files (PDF/DOCX) and returns JSON format.
Uses regex patterns for initial extraction and OpenAI as fallback for structured data.
"""

import re
import json
import io
from typing import Dict, List, Optional, BinaryIO
import pdfplumber
from docx import Document
from openai import OpenAI
import os
from datetime import datetime


class ResumeParser:
    """Parse resumes from PDF or DOCX files and extract structured information."""
    
    def __init__(self, openai_api_key: Optional[str] = None):
        """
        Initialize the resume parser.
        
        Args:
            openai_api_key: OpenAI API key. If not provided, will try to get from environment.
        """
        self.openai_api_key = openai_api_key or os.getenv("OPENAI_API_KEY")
        if self.openai_api_key:
            self.client = OpenAI(api_key=self.openai_api_key)
        else:
            self.client = None
    
    def parse_resume(self, file_content: bytes, file_type: str) -> Dict:
        """
        Main method to parse resume from binary content.
        
        Args:
            file_content: Binary content of the resume file
            file_type: File type ('pdf' or 'docx')
            
        Returns:
            Dictionary with structured resume data
        """
        # Extract text from file
        text = self._extract_text(file_content, file_type)
        
        # Try to extract structured data using regex patterns
        structured_data = self._extract_structured_data(text)
        
        # If OpenAI is available and data is incomplete, use it as fallback
        if self.client and self._is_data_incomplete(structured_data):
            structured_data = self._extract_with_openai(text, structured_data)
        
        return structured_data
    
    def _extract_text(self, file_content: bytes, file_type: str) -> str:
        """
        Extract text from PDF or DOCX file.
        
        Args:
            file_content: Binary content of the file
            file_type: Type of file ('pdf' or 'docx')
            
        Returns:
            Extracted text as string
        """
        text = ""
        
        try:
            if file_type.lower() in ['pdf', 'application/pdf']:
                text = self._extract_from_pdf(file_content)
            elif file_type.lower() in ['docx', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
                text = self._extract_from_docx(file_content)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        except Exception as e:
            raise Exception(f"Error extracting text from file: {str(e)}")
        
        return text
    
    def _extract_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF using pdfplumber."""
        text = ""
        with pdfplumber.open(io.BytesIO(file_content)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text
    
    def _extract_from_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX using python-docx."""
        text = ""
        doc = Document(io.BytesIO(file_content))
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        return text
    
    def _extract_structured_data(self, text: str) -> Dict:
        """
        Extract structured data using regex patterns.
        
        Args:
            text: Extracted text from resume
            
        Returns:
            Dictionary with extracted fields
        """
        data = {
            "name": self._extract_name(text),
            "email": self._extract_email(text),
            "phone": self._extract_phone(text),
            "skills": self._extract_skills(text),
            "experience_years": self._extract_experience_years(text),
            "education": self._extract_education(text),
            "summary": self._extract_summary(text)
        }
        
        return data
    
    def _extract_name(self, text: str) -> Optional[str]:
        """Extract name from resume text (usually first line or top of document)."""
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        
        if not lines:
            return None
        
        # Name is typically in the first 3 lines and is all caps or title case
        for line in lines[:3]:
            # Skip lines that look like email, phone, or common headers
            if any(keyword in line.lower() for keyword in ['email', 'phone', 'address', 'linkedin', 'github', '@']):
                continue
            
            # Check if line looks like a name (2-4 words, mostly letters)
            words = line.split()
            if 2 <= len(words) <= 4 and all(word.replace('.', '').isalpha() for word in words):
                return line
        
        # Fallback to first non-empty line
        return lines[0] if lines else None
    
    def _extract_email(self, text: str) -> Optional[str]:
        """Extract email address from text."""
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        matches = re.findall(email_pattern, text)
        return matches[0] if matches else None
    
    def _extract_phone(self, text: str) -> Optional[str]:
        """Extract phone number from text."""
        # Various phone number patterns
        phone_patterns = [
            r'\+?\d{1,3}[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',  # +1-234-567-8900
            r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',  # (234) 567-8900
            r'\d{3}[-.\s]?\d{3}[-.\s]?\d{4}',  # 234-567-8900
            r'\d{10}',  # 2345678900
        ]
        
        for pattern in phone_patterns:
            matches = re.findall(pattern, text)
            if matches:
                return matches[0]
        
        return None
    
    def _extract_skills(self, text: str) -> List[str]:
        """Extract skills from resume text."""
        skills = []
        
        # Common technical skills to look for
        common_skills = [
            'python', 'java', 'javascript', 'c++', 'c#', 'ruby', 'php', 'swift', 'kotlin',
            'react', 'angular', 'vue', 'node.js', 'django', 'flask', 'spring', 'express',
            'sql', 'mysql', 'postgresql', 'mongodb', 'redis', 'elasticsearch',
            'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'jenkins', 'git',
            'machine learning', 'deep learning', 'nlp', 'computer vision',
            'tensorflow', 'pytorch', 'scikit-learn', 'pandas', 'numpy',
            'html', 'css', 'typescript', 'rest api', 'graphql',
            'agile', 'scrum', 'jira', 'confluence'
        ]
        
        text_lower = text.lower()
        
        # Find skills section
        skills_section = self._extract_section(text, ['skills', 'technical skills', 'expertise', 'technologies'])
        
        if skills_section:
            search_text = skills_section.lower()
        else:
            search_text = text_lower
        
        # Look for common skills
        for skill in common_skills:
            if skill in search_text:
                skills.append(skill.title())
        
        # Remove duplicates and sort
        skills = sorted(list(set(skills)))
        
        return skills
    
    def _extract_experience_years(self, text: str) -> Optional[float]:
        """Extract years of experience from resume."""
        experience_patterns = [
            r'(\d+)\+?\s*years?\s+(?:of\s+)?experience',
            r'experience[:\s]+(\d+)\+?\s*years?',
            r'(\d+)\+?\s*years?\s+in\s+',
        ]
        
        for pattern in experience_patterns:
            matches = re.findall(pattern, text.lower())
            if matches:
                return float(matches[0])
        
        # Try to calculate from work history dates
        date_ranges = re.findall(r'(\d{4})\s*[-–—]\s*(\d{4}|present|current)', text.lower())
        
        if date_ranges:
            total_years = 0
            current_year = datetime.now().year
            
            for start, end in date_ranges:
                start_year = int(start)
                end_year = current_year if end in ['present', 'current'] else int(end)
                total_years += (end_year - start_year)
            
            return float(total_years) if total_years > 0 else None
        
        return None
    
    def _extract_education(self, text: str) -> List[Dict[str, str]]:
        """Extract education information from resume."""
        education = []
        
        # Find education section
        education_section = self._extract_section(text, ['education', 'academic', 'qualification'])
        
        if not education_section:
            return education
        
        # Common degree patterns
        degree_patterns = [
            r'(Bachelor|B\.?S\.?|B\.?A\.?|Master|M\.?S\.?|M\.?A\.?|MBA|Ph\.?D\.?|Doctorate)[^\n]*',
            r'(Associate|Diploma|Certificate)[^\n]*'
        ]
        
        lines = education_section.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            for pattern in degree_patterns:
                match = re.search(pattern, line, re.IGNORECASE)
                if match:
                    # Extract year if present
                    year_match = re.search(r'\b(19|20)\d{2}\b', line)
                    year = year_match.group(0) if year_match else None
                    
                    education.append({
                        "degree": match.group(0),
                        "year": year,
                        "details": line
                    })
                    break
        
        return education
    
    def _extract_summary(self, text: str) -> Optional[str]:
        """Extract professional summary or objective from resume."""
        summary_keywords = ['summary', 'objective', 'profile', 'about', 'overview']
        
        summary_section = self._extract_section(text, summary_keywords)
        
        if summary_section:
            # Get first few sentences (up to 500 characters)
            sentences = summary_section.split('.')[:3]
            summary = '. '.join(sentences).strip()
            
            if len(summary) > 500:
                summary = summary[:500] + '...'
            
            return summary if summary else None
        
        # Fallback: get first paragraph that's not a name/contact
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        for i, line in enumerate(lines):
            if len(line) > 100 and not any(keyword in line.lower() for keyword in ['email', 'phone', 'address', '@']):
                return line[:500]
        
        return None
    
    def _extract_section(self, text: str, keywords: List[str]) -> Optional[str]:
        """Extract a specific section from resume text based on keywords."""
        lines = text.split('\n')
        section_lines = []
        in_section = False
        
        for i, line in enumerate(lines):
            line_lower = line.lower().strip()
            
            # Check if this line is a section header
            if any(keyword in line_lower for keyword in keywords):
                in_section = True
                continue
            
            # Check if we've reached a new section
            if in_section:
                # Common section headers that indicate end of current section
                other_sections = ['experience', 'work', 'employment', 'projects', 'certifications', 
                                'awards', 'publications', 'references', 'skills', 'education']
                
                if any(section in line_lower for section in other_sections) and \
                   not any(keyword in line_lower for keyword in keywords):
                    break
                
                section_lines.append(line)
        
        return '\n'.join(section_lines).strip() if section_lines else None
    
    def _is_data_incomplete(self, data: Dict) -> bool:
        """Check if extracted data is incomplete and needs OpenAI fallback."""
        # Check if critical fields are missing
        critical_fields = ['name', 'email']
        
        for field in critical_fields:
            if not data.get(field):
                return True
        
        # Check if skills or education are empty
        if not data.get('skills') or not data.get('education'):
            return True
        
        return False
    
    def _extract_with_openai(self, text: str, existing_data: Dict) -> Dict:
        """
        Use OpenAI to extract structured data from resume text.
        
        Args:
            text: Resume text
            existing_data: Already extracted data from regex patterns
            
        Returns:
            Enhanced data dictionary
        """
        if not self.client:
            return existing_data
        
        # Define the JSON schema for extraction
        schema = {
            "name": "string (full name of the candidate)",
            "email": "string (email address)",
            "phone": "string (phone number)",
            "skills": "array of strings (technical and professional skills)",
            "experience_years": "number (total years of experience)",
            "education": "array of objects with degree, institution, year, and field",
            "summary": "string (professional summary or objective)"
        }
        
        prompt = f"""
Extract structured information from the following resume text and return it as a valid JSON object.

JSON Schema:
{json.dumps(schema, indent=2)}

Rules:
1. Extract all available information accurately
2. For skills, include both technical and soft skills
3. Calculate experience_years from work history dates if not explicitly stated
4. For education, include degree, institution, year, and field of study
5. Return ONLY valid JSON, no additional text
6. If a field is not found, use null for strings/numbers or empty array [] for arrays

Resume Text:
{text[:4000]}

Return the JSON object:
"""
        
        try:
            response = self.client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {
                        "role": "system",
                        "content": "You are an expert resume parser. Extract structured information and return valid JSON only."
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                temperature=0.3,
                max_tokens=1500
            )
            
            # Parse OpenAI response
            ai_response = response.choices[0].message.content.strip()
            
            # Remove markdown code blocks if present
            if ai_response.startswith('```'):
                ai_response = ai_response.split('```')[1]
                if ai_response.startswith('json'):
                    ai_response = ai_response[4:]
                ai_response = ai_response.strip()
            
            ai_data = json.loads(ai_response)
            
            # Merge with existing data, preferring OpenAI data when available
            for key in ['name', 'email', 'phone', 'summary', 'experience_years']:
                if ai_data.get(key) and not existing_data.get(key):
                    existing_data[key] = ai_data[key]
            
            # For arrays, merge and deduplicate
            if ai_data.get('skills'):
                existing_skills = set(existing_data.get('skills', []))
                ai_skills = set(ai_data['skills'])
                existing_data['skills'] = sorted(list(existing_skills | ai_skills))
            
            if ai_data.get('education'):
                existing_data['education'] = ai_data['education']
            
            return existing_data
            
        except Exception as e:
            print(f"OpenAI extraction failed: {str(e)}")
            return existing_data
    
    def parse_resume_from_file(self, file_path: str) -> Dict:
        """
        Convenience method to parse resume from file path.
        
        Args:
            file_path: Path to resume file
            
        Returns:
            Dictionary with structured resume data
        """
        with open(file_path, 'rb') as f:
            file_content = f.read()
        
        file_type = 'pdf' if file_path.endswith('.pdf') else 'docx'
        return self.parse_resume(file_content, file_type)


def parse_resume(file_content: bytes, file_type: str, openai_api_key: Optional[str] = None) -> Dict:
    """
    Convenience function to parse a resume.
    
    Args:
        file_content: Binary content of the resume file
        file_type: File type ('pdf' or 'docx')
        openai_api_key: Optional OpenAI API key
        
    Returns:
        Dictionary with structured resume data
    """
    parser = ResumeParser(openai_api_key=openai_api_key)
    return parser.parse_resume(file_content, file_type)
